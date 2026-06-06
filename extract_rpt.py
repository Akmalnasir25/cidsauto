import json, re, os
from docx import Document

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(BASE, 'data')
os.makedirs(OUT, exist_ok=True)

MINGGU_TARIKH = {
    1: ('12/01/2026','16/01/2026'), 2: ('19/01/2026','23/01/2026'),
    3: ('26/01/2026','30/01/2026'), 4: ('02/02/2026','06/02/2026'),
    5: ('09/02/2026','13/02/2026'), 6: ('23/02/2026','27/02/2026'),
    7: ('02/03/2026','06/03/2026'), 8: ('09/03/2026','13/03/2026'),
    9: ('16/03/2026','20/03/2026'), 10:('30/03/2026','03/04/2026'),
    11:('06/04/2026','10/04/2026'), 12:('13/04/2026','17/04/2026'),
    13:('20/04/2026','24/04/2026'), 14:('27/04/2026','01/05/2026'),
    15:('04/05/2026','08/05/2026'), 16:('11/05/2026','15/05/2026'),
    17:('18/05/2026','22/05/2026'), 18:('08/06/2026','12/06/2026'),
    19:('15/06/2026','19/06/2026'), 20:('22/06/2026','26/06/2026'),
    21:('29/06/2026','03/07/2026'), 22:('06/07/2026','10/07/2026'),
    23:('13/07/2026','17/07/2026'), 24:('20/07/2026','24/07/2026'),
    25:('27/07/2026','31/07/2026'), 26:('03/08/2026','07/08/2026'),
    27:('10/08/2026','14/08/2026'), 28:('17/08/2026','21/08/2026'),
    29:('24/08/2026','28/08/2026'), 30:('07/09/2026','11/09/2026'),
    31:('14/09/2026','18/09/2026'), 32:('21/09/2026','25/09/2026'),
    33:('28/09/2026','02/10/2026'), 34:('05/10/2026','09/10/2026'),
    35:('12/10/2026','16/10/2026'), 36:('19/10/2026','23/10/2026'),
    37:('26/10/2026','30/10/2026'), 38:('02/11/2026','06/11/2026'),
    39:('09/11/2026','13/11/2026'), 40:('16/11/2026','20/11/2026'),
    41:('23/11/2026','27/11/2026'), 42:('30/11/2026','04/12/2026'),
}

def dedup_cells(row):
    unique_cells = []
    prev_tc = None
    for cell in row.cells:
        tc = cell._tc
        if tc is not prev_tc:
            unique_cells.append(cell.text.strip())
            prev_tc = tc
        else:
            unique_cells.append(None)
    return unique_cells

def get_minggu(text):
    t = text.strip()
    if not t or 'CUTI' in t.upper() or 'ORIENTASI' in t.upper() or 'UASA' in t.upper() or 'AKHIR' in t.upper():
        return []
    nums = re.findall(r'\b(\d{1,2})\b', t.split('\n')[0].split('Kump')[0])
    return [int(n) for n in nums if 1 <= int(n) <= 42]

def process_pj(filepath):
    doc = Document(filepath)
    result = {}
    for table in doc.tables:
        for row in table.rows[1:]:
            cells = dedup_cells(row)
            ncols = sum(1 for c in cells if c is not None)
            non_none = [c for c in cells if c is not None]
            if len(non_none) < 4: continue
            if len(non_none) == 5:
                mtxt, ttxt, sk, sp = non_none[0], non_none[1], non_none[2], non_none[3]
            elif len(non_none) == 6:
                mtxt = non_none[0]
                ttxt = non_none[2] if non_none[1] == non_none[2] else non_none[1]
                sk, sp = non_none[3], non_none[4]
            else: continue
            all_txt = ' '.join(non_none)
            weeks = get_minggu(mtxt)
            if not weeks: continue
            cuti = 'CUTI' in all_txt.upper()
            uasa = 'UASA' in all_txt.upper()
            unit, tajuk = parse_bidang_tajuk(ttxt)
            for w in weeks:
                if w not in result:
                    result[w] = {
                        'minggu': w,
                        'bidang': unit,
                        'tajuk': tajuk,
                        'standard_kandungan': sk.replace('\n', ' '),
                        'standard_pembelajaran': sp.replace('\n', ' '),
                        'is_cuti': cuti,
                        'is_uasa': uasa,
                    }
    return result

def parse_bidang_tajuk(text):
    t = text.strip()
    unit = tajuk = ''
    if 'TAJUK:' in t:
        idx = t.index('TAJUK:')
        unit = t[:idx].strip()
        tajuk = t[idx+6:].strip()
    elif 'UNIT' in t:
        unit = t
    return unit, tajuk

def process_mat(filepath):
    doc = Document(filepath)
    result = {}
    for table in doc.tables:
        rows = table.rows
        header_dedup = dedup_cells(rows[0])
        non_none = [c for c in header_dedup if c is not None]
        bidang = ''
        tajuk = ''
        for cell_text in non_none:
            if cell_text.startswith('BIDANG PEMBELAJARAN:'):
                bidang = cell_text.split('BIDANG PEMBELAJARAN:')[1].strip().replace('\n', ' ')
            elif cell_text.startswith('TAJUK:'):
                tajuk = cell_text.split('TAJUK:')[1].strip().replace('\n', ' ')

        for row in rows[1:]:
            cells = dedup_cells(row)
            non_none = [c for c in cells if c is not None]
            if len(non_none) < 2: continue
            mtxt = non_none[0]
            if 'MINGGU' not in mtxt.upper() and not re.search(r'\d', mtxt): continue
            if any(x in mtxt.upper() for x in ['CUTI','ORIENTASI','UASA','AKHIR','KESEDIAAN']): continue
            sp_text = non_none[1].replace('\n', ' ') if non_none[1] else ''
            if not sp_text.strip(): continue

            nums = re.findall(r'(\d+)\b', mtxt)
            weeks = [int(n) for n in nums if 1 <= int(n) <= 42]
            if len(weeks) > 5: weeks = []
            if not weeks:
                range_m = re.findall(r'MINGGU:\s*(\d+)', mtxt)
                weeks = [int(m) for m in range_m]
            if not weeks:
                range_m = re.findall(r'MINGGU\s+(\d+)', mtxt)
                weeks = [int(m) for m in range_m]
            ranges = re.findall(r'(\d+)\s*(?:-|HINGGA|HINGGA)\s*(\d+)', mtxt)
            for s, e in ranges:
                for w in range(int(s), int(e)+1):
                    if w not in weeks and 1 <= w <= 42:
                        weeks.append(w)
            for w in weeks:
                if w not in result:
                    result[w] = {
                        'minggu': w,
                        'bidang': bidang.replace('\n', ' '),
                        'tajuk': tajuk.replace('\n', ' '),
                        'standard_kandungan': sp_text,
                        'standard_pembelajaran': sp_text,
                        'is_cuti': False,
                        'is_uasa': False,
                    }
    return result

def get_minggu_rbt(text):
    cleaned = re.sub(r'\d{1,2}[./:]\d{1,2}[./:]\d{2,4}', '', text)
    cleaned = re.sub(r'Kump\s*[A-Z]:\s*', '', cleaned)
    cleaned = re.sub(r'\d{4}', '', cleaned)
    nums = re.findall(r'\b(\d{1,2})\b', cleaned)
    return list(set([int(n) for n in nums if 2 <= int(n) <= 37]))

def process_rbt(filepath):
    doc = Document(filepath)
    result = {}
    for table in doc.tables:
        row0 = table.rows[0]
        modul_name = row0.cells[0].text.strip()
        if modul_name.startswith('MODUL:'):
            modul_name = modul_name[6:].strip()
        if modul_name.startswith('MODUL'):
            modul_name = modul_name[5:].strip()

        for row in table.rows[1:]:
            cells = dedup_cells(row)
            non_none = [c for c in cells if c is not None]
            if len(non_none) < 3: continue
            header = cells[0]
            if header and ('STANDARD KANDUNGAN' in header.upper() or 'STANDARD PEMBELAJARAN' in header.upper() or 'TAHAP PENGUASAAN' in header.upper() or 'STANDARD PRESTASI' in header.upper()):
                continue

            mtxt = non_none[0]
            if any(x in mtxt.upper() for x in ['CUTI','ORIENTASI','UASA','AKHIR','KESEDIAAN']): continue
            if 'HEADER' in mtxt.upper() or 'SUBHEADER' in mtxt.upper(): continue
            if re.match(r'^(BIDANG|STANDARD|CATATAN)', mtxt.upper()): continue

            weeks = get_minggu_rbt(mtxt)
            if not weeks: continue

            sk = (non_none[1] or '').replace('\n', ' ').strip()
            sp_text = (non_none[2] or '').replace('\n', ' ').strip() if len(non_none) > 2 else ''

            for w in sorted(weeks):
                if w not in result:
                    result[w] = {
                        'minggu': w,
                        'bidang': modul_name.replace('\n', ' ').strip(),
                        'tajuk': '',
                        'standard_kandungan': sk,
                        'standard_pembelajaran': sp_text,
                        'is_cuti': False,
                        'is_uasa': False,
                    }
    return result

files = {
    'pj6': (r'rpt_pj6.docx', r'RPT PJ TAHUN 6 2026 EDIT.docx', process_pj),
    'pj5': (r'rpt_pj5.docx', r'RPT_PJ_THN_5_SK_2026_By_Rozayus_Academy_Kump_A_&_Kump_B.docx', process_pj),
    'pj4': (r'rpt_pj4.docx', r'RPT-PJ-THN-4-SK-2025-2026-By-Rozayus-Academy-Kump-B.docx', process_pj),
    'mat6': (r'rpt_mat6.docx', r'TAHUN 6 RPT MATEMATIK 2026.docx', process_mat),
    'rbt6': (r'rpt_rbt6.docx', r'RPT_RBT_THN_6_SK_2026_By_Rozayus_Academy_Kump_A_&_Kump_B.docx', process_rbt),
    'rbt5': (r'rpt_rbt5.docx', r'RPT_RBT_THN_5_SK_2026_By_Rozayus_Academy_Kump_A_&_Kump_B.docx', process_rbt),
}

all_data = {}
# Support multi-profile: read marker file for which rpt file to use
marker_path = os.path.join(BASE, '_active_rpt_file.txt')
rpt_filename = 'rpt-data.json'  # default
if os.path.exists(marker_path):
    with open(marker_path, 'r', encoding='utf-8') as f:
        rpt_filename = f.read().strip() or rpt_filename
    print(f'Using active rpt file: {rpt_filename}')

outpath = os.path.join(OUT, rpt_filename)
if os.path.exists(outpath):
    with open(outpath, 'r', encoding='utf-8') as f:
        all_data = json.load(f)

for key, fnames_and_func in files.items():
    func = fnames_and_func[-1]
    fnames = fnames_and_func[:-1]
    
    fpath = None
    for fname in fnames:
        p = os.path.join(BASE, fname)
        if os.path.exists(p):
            fpath = p
            break
            
    if not fpath:
        continue
        
    print(f'Processing {key}: {os.path.basename(fpath)}')
    try:
        data = func(fpath)
        all_data[key] = data
        print(f'  Found {len(data)} weeks')
        for w in sorted(data.keys(), key=int)[:3]:
            d = data[w]
            print(f'  Minggu {w}: bidang={d.get("bidang","")[:40]} | tajuk={d.get("tajuk","")[:40]}')
    except Exception as e:
        print(f'  ERROR: {e}')
        if key not in all_data:
            all_data[key] = {}

all_data['minggu_tarikh'] = {str(k): list(v) for k, v in MINGGU_TARIKH.items()}

print(f'\nSaving merged data to {outpath}')
with open(outpath, 'w', encoding='utf-8') as f:
    json.dump(all_data, f, ensure_ascii=False, indent=2)
print(f'Saved to {outpath}')
print(f'Total keys: {list(all_data.keys())}')
